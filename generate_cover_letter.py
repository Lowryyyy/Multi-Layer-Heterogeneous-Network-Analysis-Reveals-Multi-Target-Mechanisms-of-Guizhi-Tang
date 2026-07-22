"""Generate Cover Letter for Journal of Ethnopharmacology submission."""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()
style = doc.styles['Normal']
style.font.name = "Times New Roman"
style.font.size = Pt(12)

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# Date
p = doc.add_paragraph()
p.add_run("July 21, 2026").font.size = Pt(12)

# Editor
p = doc.add_paragraph()
p.add_run("Dear Editor-in-Chief,").font.size = Pt(12)
doc.add_paragraph()

# Title
p = doc.add_paragraph()
run = p.add_run("Re: Submission of manuscript entitled ")
run.font.size = Pt(12)
run = p.add_run('"Drug Repurposing Prediction for Guizhi Tang via Multi-Layer Heterogeneous Knowledge Graph and Graph Neural Networks"')
run.font.size = Pt(12)
run.font.italic = True

doc.add_paragraph()

# Body
body_paras = [
    "We are pleased to submit our original research article for consideration for publication in the Journal of Ethnopharmacology. This manuscript presents a computational framework that integrates network pharmacology with heterogeneous graph neural networks (GNNs) to predict novel drug repurposing opportunities for Guizhi Tang, a classical Traditional Chinese Medicine formula with over 1,800 years of clinical use.",

    "Guizhi Tang, composed of five herbs (Ramulus Cinnamomi, Radix Paeoniae Alba, Rhizoma Zingiberis Recens, Fructus Jujubae, and Radix Glycyrrhizae), has been traditionally used to harmonize Ying and Wei levels and release the exterior. While modern pharmacological studies have confirmed its anti-inflammatory, immunomodulatory, and cardiovascular protective activities, its broader therapeutic potential remains incompletely explored. Conventional network pharmacology approaches rely on static enrichment analysis and cannot capture the complex, non-linear relationships embedded in multi-layer pharmacological networks.",

    "In this study, we constructed a heterogeneous knowledge graph comprising 189 nodes (5 herbs, 77 compounds, 47 targets, 31 diseases, 29 drugs) and 1,602 edges across seven relation types, integrating data from TCMSP, DrugBank, DisGeNET, OMIM, and STRING databases. We systematically compared five link prediction models and demonstrated that a Heterogeneous Graph Neural Network with multi-head attention (HetGNN) achieved the best balanced performance (AUC-ROC 0.8707, AUPRC 0.6133, threshold-optimized F1 0.560), outperforming MLP baseline, homogeneous GCN, and R-GCN. Five-fold cross-validation confirmed model robustness (AUC 0.8936 \u00b1 0.032).",

    "The expanded knowledge graph yielded diverse, biologically meaningful predictions. Kaempferol emerged as the top-ranked compound for pancreatic cancer (prediction score 0.8455, 12 shared targets), while 6-gingerol was predicted for breast cancer (0.8063) and lung cancer (0.7877). Molecular docking validation confirmed favorable binding affinities for all 26 compound-target pairs (mean binding energy \u22126.80 to \u22127.29 kcal/mol). KEGG pathway enrichment analysis identified significant involvement of cancer pathways, MAPK signaling, apoptosis, PI3K-Akt signaling, and NF-\u03baB signaling.",

    "We believe this work is well suited for the Journal of Ethnopharmacology because it: (1) addresses the pharmacological potential of a classical TCM formula using state-of-the-art computational methods; (2) provides a systematic, reproducible framework for TCM drug repurposing that integrates traditional knowledge with modern pharmacological data; (3) generates experimentally testable hypotheses supported by molecular docking and pathway enrichment validation; and (4) demonstrates the value of heterogeneous graph architectures for capturing the multi-component, multi-target nature of TCM formulas.",

    "This manuscript has not been published previously and is not under consideration elsewhere. All authors have read and approved the submitted version. The authors declare no competing interests.",

    "We hope that the reviewers will find our work interesting and suitable for publication in your esteemed journal. We look forward to your favorable response.",
]

for text in body_paras:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.size = Pt(12)

doc.add_paragraph()

# Closing
p = doc.add_paragraph()
p.add_run("Sincerely,").font.size = Pt(12)
doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("Du Yazhou")
run.font.size = Pt(12)
run.font.bold = True

p = doc.add_paragraph()
p.add_run("Xinjiang Second Medical College").font.size = Pt(11)

p = doc.add_paragraph()
p.add_run("Karamay 834000, Xinjiang, China").font.size = Pt(11)

p = doc.add_paragraph()
p.add_run("E-mail: 861444938@qq.com").font.size = Pt(11)

# Save
output_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "results")
os.makedirs(output_dir, exist_ok=True)
doc.save(os.path.join(output_dir, "Cover_Letter_JEP.docx"))
print("Cover Letter saved.")
