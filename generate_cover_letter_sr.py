"""Generate Cover Letter for Scientific Reports submission."""
from docx import Document
from docx.shared import Pt, Cm
import os

doc = Document()
style = doc.styles['Normal']
style.font.name = "Times New Roman"
style.font.size = Pt(12)
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

p = doc.add_paragraph()
p.add_run("July 22, 2026").font.size = Pt(12)
p = doc.add_paragraph()
p.add_run("Dear Editor,").font.size = Pt(12)
doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("Re: Submission of manuscript entitled ")
run.font.size = Pt(12)
run = p.add_run('"Multi-Layer Heterogeneous Network Analysis Reveals Multi-Target Mechanisms of Guizhi Tang: Integrating Graph Neural Networks with Molecular Docking and Pathway Enrichment"')
run.font.size = Pt(12)
run.font.italic = True
doc.add_paragraph()

body = [
    "We are pleased to submit our original research article for consideration for publication in Scientific Reports. This manuscript presents an integrated computational framework that combines heterogeneous graph neural networks, molecular docking, and pathway enrichment analysis to systematically elucidate the multi-target molecular mechanisms of Guizhi Tang, a classical Traditional Chinese Medicine formula with over 1,800 years of clinical use.",

    "Guizhi Tang, composed of five herbs, has been traditionally used to treat common cold, rheumatic disease, and gastric disorders. While its anti-inflammatory and immunomodulatory activities have been confirmed in modern pharmacological studies, the molecular mechanisms underlying its therapeutic effects\u2014and its potential applications beyond traditional indications\u2014remain incompletely understood. Conventional network pharmacology approaches rely on static enrichment analysis and cannot capture the complex, non-linear relationships in multi-layer pharmacological networks.",

    "In this study, we constructed a heterogeneous knowledge graph (189 nodes, 1,602 edges) integrating data from five public databases and employed a Heterogeneous Graph Neural Network (HetGNN) to predict compound-disease associations. The model achieved AUC-ROC of 0.8707 (five-fold CV: 0.8936 \u00b1 0.032). The top predictions identified kaempferol as a key compound for pancreatic cancer (12 shared targets, score 0.8455) and 6-gingerol for breast and lung cancers. Molecular docking confirmed favorable binding for all 26 compound-target pairs (mean energy \u22126.80 to \u22127.29 kcal/mol). KEGG enrichment revealed significant involvement of cancer pathways (p = 1.2 \u00d7 10\u207b\u2078), MAPK signaling, apoptosis, PI3K-Akt, and NF-\u03baB pathways. We further demonstrated concordance between traditional anti-inflammatory indications and the TNF/NF-\u03baB/IL6 pathways identified computationally.",

    "We believe this work is well suited for Scientific Reports because it: (1) presents a methodologically rigorous computational framework with comprehensive model comparison, cross-validation, and ablation analysis; (2) provides structural-level validation through molecular docking; (3) identifies enriched biological pathways connecting traditional therapeutic uses to modern molecular understanding; and (4) is fully reproducible, with all code, data, and configuration files made publicly available.",

    "This manuscript has not been published previously and is not under consideration elsewhere. All authors have read and approved the submitted version. The authors declare no competing interests.",
]

for text in body:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    p.add_run(text).font.size = Pt(12)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Sincerely,").font.size = Pt(12)
doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("Du Yazhou")
r.font.size = Pt(12)
r.font.bold = True
p = doc.add_paragraph()
p.add_run("Xinjiang Second Medical College, Karamay 834000, Xinjiang, China").font.size = Pt(11)
p = doc.add_paragraph()
p.add_run("E-mail: 861444938@qq.com").font.size = Pt(11)

output_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "results")
doc.save(os.path.join(output_dir, "Cover_Letter_SR.docx"))
print("Cover Letter for Scientific Reports saved.")
