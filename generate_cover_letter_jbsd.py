"""Generate Cover Letter for JBSD submission."""
from docx import Document
from docx.shared import Pt, Cm
import os

doc = Document()
style = doc.styles['Normal']
style.font.name = "Times New Roman"
style.font.size = Pt(12)
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

p = doc.add_paragraph()
p.add_run("July 22, 2026").font.size = Pt(12)

p = doc.add_paragraph()
p.add_run("Dear Editor-in-Chief,").font.size = Pt(12)
doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("Re: Submission of manuscript entitled ")
run.font.size = Pt(12)
run = p.add_run('"Multi-Layer Heterogeneous Network Analysis Reveals Multi-Target Mechanisms of Guizhi Tang: Integrating Graph Neural Networks with Molecular Docking and Pathway Enrichment"')
run.font.size = Pt(12)
run.font.italic = True
doc.add_paragraph()

body = [
    "We are pleased to submit our original research article for consideration for publication in the Journal of Biomolecular Structure and Dynamics. This manuscript presents a computational framework that integrates heterogeneous graph neural networks, molecular docking, and pathway enrichment analysis to systematically elucidate the multi-target molecular mechanisms of Guizhi Tang, a classical Traditional Chinese Medicine formula with over 1,800 years of clinical use.",

    "Guizhi Tang, composed of five herbs (Ramulus Cinnamomi, Radix Paeoniae Alba, Rhizoma Zingiberis Recens, Fructus Jujubae, and Radix Glycyrrhizae), has been traditionally used to treat common cold, rheumatic disease, and gastric disorders. While modern pharmacological studies have confirmed its anti-inflammatory and immunomodulatory activities, the molecular mechanisms underlying its therapeutic effects remain incompletely understood, particularly regarding its potential applications beyond traditional indications.",

    "In this study, we constructed a heterogeneous knowledge graph (189 nodes, 1,602 edges) integrating data from TCMSP, DrugBank, DisGeNET, OMIM, and STRING, and employed a Heterogeneous Graph Neural Network (HetGNN) to predict compound-disease associations. The model achieved AUC-ROC of 0.8707 and AUPRC of 0.6133. The top predictions identified kaempferol as a key compound for pancreatic cancer (12 shared targets, prediction score 0.8455) and 6-gingerol for breast cancer (0.8063) and lung cancer (0.7877). Molecular docking validation confirmed favorable binding affinities for all 26 compound-target pairs (mean binding energy -6.80 to -7.29 kcal/mol). KEGG pathway enrichment revealed significant involvement of Pathways in cancer (p = 1.2 \u00d7 10\u207b\u2078), MAPK signaling, apoptosis, PI3K-Akt signaling, and NF-\u03baB signaling. We further demonstrated concordance between Guizhi Tang's traditional anti-inflammatory indications and the TNF/NF-\u03baB/IL6 pathways identified through our computational analysis.",

    "We believe this work is well suited for the Journal of Biomolecular Structure and Dynamics because it: (1) elucidates the multi-target molecular mechanisms of a traditional medicine formula using state-of-the-art computational methods; (2) provides structural-level validation through molecular docking of predicted compound-target interactions; (3) identifies enriched biological pathways that connect traditional therapeutic uses to modern molecular understanding; and (4) demonstrates the value of heterogeneous graph architectures for capturing the multi-component pharmacology inherent to traditional medicine formulas.",

    "This manuscript has not been published previously and is not under consideration elsewhere. All authors have read and approved the submitted version. The authors declare no competing interests.",
]

for text in body:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    p.add_run(text).font.size = Pt(12)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Sincerely,").font.size = Pt(12)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Du Yazhou").font.size = Pt(12)
p.runs[0].font.bold = True
p = doc.add_paragraph()
p.add_run("Xinjiang Second Medical College, Karamay 834000, Xinjiang, China").font.size = Pt(11)
p = doc.add_paragraph()
p.add_run("E-mail: 861444938@qq.com").font.size = Pt(11)

output_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "results")
doc.save(os.path.join(output_dir, "Cover_Letter_JBSD.docx"))
print("Cover Letter saved.")
